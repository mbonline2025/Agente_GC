from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from app.auth import verify_token
import pandas as pd
import os
from docx import Document
from fpdf import FPDF
import tempfile
from datetime import datetime

router = APIRouter()

@router.get("/ask/exportar", tags=["Exportar"])
def exportar_resposta(formato: str = Query(..., enum=["pdf", "doc", "excel"]), user=Depends(verify_token)):
    historico_path = "logs/historico_conversas.csv"

    if not os.path.exists(historico_path):
        raise HTTPException(status_code=404, detail="Nenhuma conversa registrada.")

    try:
        df = pd.read_csv(historico_path)
        df_user = df[df["email"] == user["email"]].sort_values(by="timestamp", ascending=False)

        if df_user.empty:
            raise HTTPException(status_code=404, detail="Nenhuma conversa encontrada para exportação.")

        ultima = df_user.iloc[0]
        pergunta = ultima["pergunta"]
        resposta = ultima["resposta"]

        if formato == "doc":
            doc = Document()
            doc.add_heading("MB.IA - Resposta Exportada", 0)
            doc.add_paragraph(f"Pergunta: {pergunta}")
            doc.add_paragraph(f"\nResposta:\n{resposta}")
            temp_path = tempfile.mktemp(suffix=".docx")
            doc.save(temp_path)
            return FileResponse(temp_path, filename="resposta.docx")

        elif formato == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, f"MB.IA - Resposta Exportada\n\nPergunta: {pergunta}\n\nResposta:\n{resposta}")
            temp_path = tempfile.mktemp(suffix=".pdf")
            pdf.output(temp_path)
            return FileResponse(temp_path, filename="resposta.pdf")

        elif formato == "excel":
            temp_path = tempfile.mktemp(suffix=".xlsx")
            export_df = pd.DataFrame([{"pergunta": pergunta, "resposta": resposta}])
            export_df.to_excel(temp_path, index=False)
            return FileResponse(temp_path, filename="resposta.xlsx")

        else:
            raise HTTPException(status_code=400, detail="Formato de exportação inválido.")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao exportar: {str(e)}")
